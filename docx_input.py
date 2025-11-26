import os
import csv
from openpyxl import load_workbook
from docx import Document
from tkinter import (
    Tk,
    Label,
    Button,
    filedialog,
    messagebox,
    StringVar,
    Frame,
    LEFT,
    RIGHT,
    BOTH,
)


def show_help():
    help_text = (
        "📌 How to use DOCX Generator:\n\n"
        "1️⃣ Select a Word template with placeholders like {{Name}}\n"
        "2️⃣ Select a CSV or XLSX file (headers = placeholder names)\n"
        "3️⃣ Click 'Generate Documents' → choose output folder\n\n"
        "💡 Tip: Field names must EXACTLY match headers (case & space sensitive!)\n\n"
        "📥 Download latest version:\n"
        "https://github.com/butterman28/Offiice_Autos/releases"
    )
    messagebox.showinfo("Help", help_text)


def replace_placeholders(doc, data):
    # Clean data: convert keys to str, handle None/NaN
    clean_data = {}
    for k, v in data.items():
        key = str(k).strip() if k is not None else ""
        # Convert None/NaN to empty string, preserve numbers/dates as str
        val = "" if v is None or (isinstance(v, float) and str(v) == "nan") else str(v)
        clean_data[key] = val

    # Replace in paragraphs
    for p in doc.paragraphs:
        for key, value in clean_data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if placeholder in inline[i].text:
                        inline[i].text = inline[i].text.replace(placeholder, value)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in clean_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)


def read_rows(data_file_path):
    """Read CSV or XLSX and return list of dict (like pandas DataFrame.to_dict('records'))"""
    if data_file_path.endswith(".csv"):
        with open(data_file_path, newline="", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            return [row for row in reader if any(row.values())]  # skip empty rows

    elif data_file_path.endswith(".xlsx"):
        wb = load_workbook(data_file_path, read_only=True, data_only=True)
        try:
            ws = wb.active
            rows = []
            # Read headers (first row)
            headers = []
            for cell in ws[1]:
                # Handle None headers (e.g., empty columns)
                h = str(cell.value).strip() if cell.value is not None else ""
                headers.append(h)

            # Read data rows
            for row in ws.iter_rows(min_row=2, values_only=True):
                if any(cell is not None for cell in row):  # skip fully empty rows
                    # Match values to headers (pad short rows)
                    values = list(row) + [""] * (len(headers) - len(row))
                    record = {headers[i]: values[i] for i in range(len(headers))}
                    rows.append(record)
            return rows
        finally:
            wb.close()
    else:
        raise ValueError("Unsupported file format. Please use CSV or XLSX.")


def generate_docs(template_path, data_file_path, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    rows = read_rows(data_file_path)

    for idx, row in enumerate(rows):
        doc = Document(template_path)
        replace_placeholders(doc, row)
        output_path = os.path.join(output_dir, f"output_{idx + 1}.docx")
        doc.save(output_path)
        print(f"Saved: {output_path}")

    messagebox.showinfo("Success", f"{len(rows)} documents created in '{output_dir}'")


def select_template():
    return filedialog.askopenfilename(
        title="Select DOCX Template", filetypes=[("Word Documents", "*.docx")]
    )


def select_data_file():
    return filedialog.askopenfilename(
        title="Select Excel or CSV File",
        filetypes=[("Spreadsheet Files", "*.xlsx *.csv")],
    )


def run_app():
    root = Tk()
    root.title("DOCX Generator")
    root.geometry("600x200")

    template_path_var = StringVar()
    data_file_path_var = StringVar()

    # Frame for template selector (left)
    frame_left = Frame(root)
    frame_left.pack(side=LEFT, fill=BOTH, expand=True, padx=10, pady=10)

    Label(frame_left, text="Select DOCX Template:").pack(pady=(0, 5))
    Button(
        frame_left,
        text="Browse Template",
        command=lambda: browse_template(template_path_var),
    ).pack()
    Label(frame_left, textvariable=template_path_var, wraplength=250).pack(pady=(5, 0))

    # Frame for data file selector (right)
    frame_right = Frame(root)
    frame_right.pack(side=RIGHT, fill=BOTH, expand=True, padx=10, pady=10)

    Label(frame_right, text="Select Excel or CSV File:").pack(pady=(0, 5))
    Button(
        frame_right,
        text="Browse Data File",
        command=lambda: browse_data_file(data_file_path_var),
    ).pack()
    Label(frame_right, textvariable=data_file_path_var, wraplength=250).pack(
        pady=(5, 0)
    )

    # Generate button at bottom center
    def on_generate():
        template_path = template_path_var.get()
        data_file_path = data_file_path_var.get()
        if not template_path:
            messagebox.showwarning(
                "Missing File", "Please select a DOCX template file."
            )
            return
        if not data_file_path:
            messagebox.showwarning(
                "Missing File", "Please select an Excel or CSV data file."
            )
            return
        output_dir = filedialog.askdirectory(title="Select Output Folder")
        if not output_dir:
            return
        try:
            generate_docs(template_path, data_file_path, output_dir)
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred:\n{e}")

    generate_btn = Button(root, text="Generate Documents", command=on_generate)
    generate_btn.pack(pady=10)
    help_btn = Button(root, text="ℹ️ Help", command=lambda: show_help())
    help_btn.pack(pady=5)

    def browse_template(var):
        path = select_template()
        if path:
            var.set(path)

    def browse_data_file(var):
        path = select_data_file()
        if path:
            var.set(path)

    root.mainloop()


if __name__ == "__main__":
    run_app()
